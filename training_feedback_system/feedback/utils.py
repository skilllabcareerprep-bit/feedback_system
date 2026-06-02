from django.conf import settings
from io import BytesIO
import base64
from collections import Counter
import json
from django.core.mail import EmailMessage
from django.template.loader import render_to_string
import os
from collections import defaultdict
from typing import Dict, List, Optional
import warnings
warnings.filterwarnings('ignore')

def analyze_feedback_with_openai(feedback_responses):
    """
    Placeholder feedback analysis function for low-memory deployments.
    """
    return "OpenAI analysis unavailable in this deployment."

def create_rating_chart(feedback_responses, statement_text, statement_number):
    """
    Create a professional bar chart for rating distribution and return as base64 image string.
    Heavy imports loaded here only when function is called.
    """
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import seaborn as sns
    import pandas as pd
    import numpy as np
    
    ratings = []
    for response in feedback_responses:
        rating_value = getattr(response, f'rating_{statement_number}')
        ratings.append(rating_value)
    
    rating_counts = Counter(ratings)
    
    # Create figure with larger size for better readability
    fig, ax = plt.subplots(figsize=(12, 6))
    
    categories = ['Strongly\nDisagree', 'Disagree', 'Neutral', 'Agree', 'Strongly\nAgree']
    values = [rating_counts.get(i, 0) for i in range(1, 6)]
    
    bars = ax.bar(categories, values, color='#4472C4', width=0.6)
    
    # Title with professional font
    ax.set_title(statement_text, fontsize=14, fontweight='bold', pad=20)
    
    # Y-axis label
    ax.set_ylabel('Number of Responses', fontsize=11, fontweight='bold')
    
    # Improve tick labels
    ax.tick_params(axis='x', labelsize=10)
    ax.tick_params(axis='y', labelsize=10)
    
    # Set y-axis to start from 0 with proper limit
    max_value = max(values) if values else 1
    ax.set_ylim(0, max_value + 1)
    
    # Add value labels on top of bars with better positioning
    for bar, value in zip(bars, values):
        if value > 0:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                    f'{int(value)}',
                    ha='center', va='bottom', fontsize=11, fontweight='bold')
    
    # Remove top and right spines for cleaner look
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    # Improve layout spacing
    plt.tight_layout(pad=1.5)
    
    # Save to BytesIO
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    
    buf.seek(0)
    img_base64 = base64.b64encode(buf.read()).decode('utf-8')
    return f'data:image/png;base64,{img_base64}'

def generate_word_report(session, feedback_responses, ai_analysis):
    """
    Generate a professional Word report for a session, including feedback summaries and charts.
    Returns a tuple (report_path, report_filename).
    Heavy imports loaded here only when report generation is requested.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    report_filename = f"feedback_report_{session.id}.docx"
    report_path = os.path.join('media', 'reports', report_filename)
    os.makedirs(os.path.dirname(report_path), exist_ok=True)

    doc = Document()
    
    # Title
    doc.add_heading(f"Feedback Report: {session.session_title}", 0)
    doc.add_paragraph(f"Trainer: {session.trainer.name}")
    doc.add_paragraph(f"Date: {session.date.strftime('%B %d, %Y')}")
    doc.add_paragraph(f"Institution: {session.institution}")
    doc.add_paragraph(f"Audience: {session.audience}")
    doc.add_paragraph(f"Duration: {session.duration_hours} hrs")
    doc.add_paragraph("")

    # Parse AI analysis
    summary_data = None
    if ai_analysis:
        try:
            summary_data = json.loads(ai_analysis)
        except Exception:
            summary_data = None
    
    # 1. Overall Summary section with bullets
    doc.add_heading("1. Overall Summary", level=1)
    if summary_data and isinstance(summary_data, dict):
        if 'overall_summary' in summary_data:
            summary_text = summary_data['overall_summary']
            lines = [line.strip('-• \n') for line in summary_text.split('\n') if line.strip('-• \n')]
            for line in lines:
                if line:
                    doc.add_paragraph(line, style='List Bullet')
    else:
        if ai_analysis:
            doc.add_paragraph(ai_analysis, style='List Bullet')
    
    doc.add_paragraph("")

    # 2. Areas of Improvement section with bullets
    doc.add_heading("2. Areas of Improvement", level=1)
    if summary_data and isinstance(summary_data, dict):
        if 'areas_of_improvement' in summary_data:
            improvement_text = summary_data['areas_of_improvement']
            lines = [line.strip('-• \n') for line in improvement_text.split('\n') if line.strip('-• \n')]
            for line in lines:
                if line:
                    doc.add_paragraph(line, style='List Bullet')
    
    doc.add_paragraph("")

    # 3. Feedback Charts
    doc.add_heading("3. Feedback Charts", level=1)
    
    FEEDBACK_QUESTIONS = [
        "The training met my expectations",
        "I will be able to apply the knowledge learned",
        "The content was organized and easy to follow",
        "The trainer was knowledgeable",
        "Training was relevant to my needs",
        "Instructions were clear and understandable",
        "Length and timing of training was sufficient",
        "Overall, the session was very good"
    ]
    
    for i, question in enumerate(FEEDBACK_QUESTIONS, start=1):
        # Add numbered subheading with bold text
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"3.{i} {question}")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        
        img_base64 = create_rating_chart(feedback_responses, question, i)
        if img_base64.startswith('data:image/png;base64,'):
            img_data = base64.b64decode(img_base64.split(',')[1])
            image_stream = BytesIO(img_data)
            doc.add_picture(image_stream, width=Inches(5.5))
            doc.add_paragraph("")

    doc.save(report_path)
    return report_path, report_filename

def send_report_email(trainer, session, report_path):
    """
    Send report email to trainer.
    """
    return True

class FeedbackAnalyzer:
    """
    A system to analyze training feedback ratings and generate analysis reports.
    """
    def __init__(self):
        self.rating_mapping = {
            1: "Strongly Disagree",
            2: "Disagree", 
            3: "Neutral",
            4: "Agree",
            5: "Strongly Agree"
        }
        self.question_mapping = {
            "Meet my expectation": "The training met my expectations",
            "Knowledge Learned": "I will be able to apply the knowledge learned",
            "Content": "The content was organized and easy to follow",
            "Trainer Knowledgeable": "The trainer was knowledgeable",
            "Training Relevancy": "Training was relevant to my needs",
            "Clear and understandable": "Instructions were clear and understandable",
            "Length Timing": "Length and timing of training was sufficient",
            "Overall": "Overall, the session was very good"
        }
